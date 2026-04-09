"""One-off generator for OTP approach Word doc; run from repo root: python docs/build_otp_doc.py"""
from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

OUT = Path(__file__).resolve().parent / "OTP-Authentication-Implementation-Approach.docx"


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)
    return p


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def main():
    doc = Document()

    t = doc.add_paragraph()
    t.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r = t.add_run("ProjectBazaar — Authentication Approach")
    r.bold = True
    r.font.size = Pt(16)

    st = doc.add_paragraph()
    st.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    sr = st.add_run("Email OTP & Phone (SMS) OTP Implementation")
    sr.font.size = Pt(14)

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.add_run("Document purpose: ").bold = True
    meta.add_run(
        "Describe how we will implement passwordless login using email one-time passwords (OTP) "
        "and SMS OTP with phone numbers, aligned with our existing AWS-centric architecture."
    )

    doc.add_paragraph()
    add_heading(doc, "1. Current technical context", level=1)
    add_para(
        doc,
        "ProjectBazaar today is a React (Vite) single-page application that calls multiple "
        "Amazon API Gateway endpoints backed by AWS Lambda. Business data and sessions are "
        "oriented around Lambda + Amazon DynamoDB (and related AWS services such as S3/SQS where applicable). "
        "Region in use includes ap-south-2. Payment flows already rely on Lambda environment "
        "variables for third-party secrets (e.g. Razorpay).",
    )

    add_heading(doc, "2. Goals", level=1)
    add_bullets(
        doc,
        [
            "Allow users to sign in (or verify identity) using a code sent to email.",
            "Allow users to sign in (or verify identity) using a code sent via SMS to a verified phone number.",
            "Minimize exposure of secrets in the browser; keep long-lived secrets only in AWS (Lambda env, Secrets Manager, or Cognito).",
            "Fit API Gateway + Lambda patterns already used across the product.",
            "Support production-scale email/SMS delivery and basic abuse controls (rate limits, expiry, attempt limits).",
        ],
    )

    add_heading(doc, "3. Recommended strategic approach", level=1)
    add_para(
        doc,
        "We recommend Amazon Cognito as the primary identity layer for both email OTP and SMS OTP, "
        "with Amazon SES for production email and Amazon SNS for SMS, integrated as Cognito requires. "
        "This gives managed OTP generation, delivery orchestration, and standard JWTs (ID/access tokens) "
        "that API Gateway can validate via a Cognito authorizer or that Lambdas can verify.",
        bold=False,
    )
    add_para(
        doc,
        "Alternative (custom) approach: implement request/verify endpoints as new Lambdas, store OTP "
        "hashes and metadata in DynamoDB with TTL, send mail via SES and SMS via SNS (or a regional SMS provider). "
        "Use this if we must avoid Cognito migration or need highly bespoke flows.",
    )

    add_heading(doc, "4. Email OTP — implementation approach", level=1)
    add_heading(doc, "4.1 Cognito path (preferred)", level=2)
    add_bullets(
        doc,
        [
            "Create or extend an Amazon Cognito User Pool with passwordless / email OTP sign-in (choice-based sign-in, ALLOW_USER_AUTH on the app client).",
            "Configure a public SPA app client (no client secret).",
            "Configure Amazon SES for outbound email in production (Cognito default limits are insufficient at scale).",
            "Frontend (React): use AWS Amplify Auth or the Cognito SDK — initiate sign-in with email, user receives code, confirm with OTP to complete authentication.",
            "Backend: attach Amazon Cognito as authorizer on API Gateway or validate JWTs in existing Lambdas; map Cognito sub/email to internal user records in DynamoDB as needed.",
        ],
    )
    add_heading(doc, "4.2 Custom Lambda path (optional)", level=2)
    add_bullets(
        doc,
        [
            "POST /auth/email-otp/send — validate email, rate-limit, generate OTP, store HMAC/hash + expiry + attempt count in DynamoDB, send email via SES.",
            "POST /auth/email-otp/verify — verify code, enforce expiry and max attempts, issue session token (signed JWT or opaque server-side session).",
            "Never store plaintext OTP in DynamoDB; enforce short TTL (e.g. 5–10 minutes) in application logic (DynamoDB TTL is eventually consistent).",
        ],
    )

    add_heading(doc, "5. Phone SMS OTP — implementation approach", level=1)
    add_heading(doc, "5.1 Cognito path (preferred)", level=2)
    add_bullets(
        doc,
        [
            "Enable SMS / phone-based passwordless OTP in the same Cognito User Pool where appropriate.",
            "Configure IAM so Cognito can publish SMS through Amazon SNS.",
            "Same client flow pattern as email: user enters E.164 phone number, receives SMS code, confirms in the app.",
            "Use Cognito’s built-in verification where possible to avoid duplicating OTP logic.",
        ],
    )
    add_heading(doc, "5.2 Regional compliance (India and similar markets)", level=2)
    add_para(
        doc,
        "For Indian mobile numbers, TRAI DLT rules apply: registered entity, approved templates, "
        "sender ID, and correct SNS message attributes may be required for reliable delivery. "
        "If Amazon SNS setup is blocked or costly for our use case, evaluate an approved third-party "
        "SMS provider (e.g. Twilio, MSG91) from Lambda while keeping verification and user records under our control.",
    )
    add_heading(doc, "5.3 Custom Lambda path (optional)", level=2)
    add_bullets(
        doc,
        [
            "Same pattern as email OTP: send and verify Lambdas, DynamoDB for hashed OTP + metadata, SNS (or provider API) for SMS.",
            "Strict rate limiting per phone number and per IP; cap verification attempts before lockout or cooldown.",
        ],
    )

    add_heading(doc, "6. Security and operations", level=1)
    add_bullets(
        doc,
        [
            "Secrets (SES/SMS credentials if any, webhook secrets) live in Lambda environment variables or AWS Secrets Manager — not in the frontend.",
            "Rate limit OTP issuance and verification endpoints to reduce abuse and cost.",
            "Log security events without logging raw OTPs or PII in clear text.",
            "Use HTTPS only; align Content-Security-Policy with any new auth domains if applicable.",
            "Document runbooks for SES/SNS sandbox exit, bounce/complaint handling, and Cognito pool recovery.",
        ],
    )

    add_heading(doc, "7. Integration with existing ProjectBazaar APIs", level=1)
    add_para(
        doc,
        "Today, authentication is centered on dedicated API Gateway + Lambda URLs consumed from the React app. "
        "After introducing Cognito (or custom OTP), new or existing Lambdas should accept authenticated identity "
        "from JWT claims (preferred) or from a server-issued session, and continue to use DynamoDB for profiles, "
        "projects, and orders. A phased rollout can keep legacy login available until users are migrated.",
    )

    add_heading(doc, "8. Summary decision", level=1)
    add_para(
        doc,
        "Primary recommendation: implement both email OTP and SMS OTP through Amazon Cognito passwordless flows, "
        "Amazon SES for email, SNS (with regional compliance) for SMS, and API Gateway JWT authorization aligned "
        "with current Lambda + DynamoDB services. Fall back to a custom Lambda + DynamoDB + SES/SNS design only if "
        "product or compliance constraints require full control over every step of the OTP lifecycle.",
    )

    doc.add_paragraph()
    foot = doc.add_paragraph()
    foot.add_run("Generated for internal planning — ProjectBazaar.").italic = True

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
